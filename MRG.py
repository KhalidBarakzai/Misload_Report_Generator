#Misload Report Generator (MRG)
#Author: Khalid Barakzai
#June, 2023
""" 
This script serves the purpose of generating a package misload report. It achieves this by executing the following tasks:
Extracting two specific columns of data from an Excel sheet and constructing a name-route ID map. This map associates Loaders with multiple route IDs.
Extracting a targeted column of data from another Excel sheet, specifically focusing on misloads.
For each occurrence of a misload associated with a Route ID, a counter is toggled to accurately capture and tally the misloads assigned to the corresponding loader.
Ultimately, the script generates a professional-looking Word document, which is saved in a designated folder named "Misload Reports".
This program has been meticulously crafted to serve as an indispensable tool for a Full-Time Preload Supervisor. """

#Run Script in Terminal:
#python3 MRG.py

##---------------------------------------------------------------------------------------------------

import os
from os.path import join
from typing import Dict, List
from collections import defaultdict
import glob
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from datetime import datetime, timedelta
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from datetime import date

##---------------------------------------------------------------------------------------------------

def find_most_recent_excel_file(folder_path: str) -> str:
    if not os.path.exists(folder_path):
        print(f"Folder not found: {folder_path}")
        return ""

    excel_files = glob.glob(join(folder_path, "*.xlsx"))
    if not excel_files:
        print(f"No Excel files found in the folder: {folder_path}")
        return ""

    sorted_files = sorted(excel_files, key=lambda x: os.path.getmtime(x), reverse=True)
    return sorted_files[0]

##---------------------------------------------------------------------------------------------------

def extract_name_route_associations(file_path: str, sheet_index: int, name_column_index1: int, route_column_index1: int) -> Dict[str, List[str]]:
    name_route_map = defaultdict(list)
    name_route_ids = defaultdict(list)
    workbook = load_workbook(file_path)
    sheet = workbook.worksheets[sheet_index]

    # Get merged cell ranges in the name column
    merged_ranges = sheet.merged_cells.ranges

    # Iterate down the rows and map route IDs to names
    for row_idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):  # Start from row 2 to skip header
        name = row[name_column_index1]
        route = row[route_column_index1]

        if name and route:
            is_merged = any(merged_range.min_row <= row_idx <= merged_range.max_row and merged_range.min_col <= name_column_index1 and merged_range.max_col >= name_column_index1 for merged_range in merged_ranges)
            if is_merged:
                name_range = None
                for merged_range in merged_ranges:
                    if merged_range.min_row <= row_idx <= merged_range.max_row and merged_range.min_col <= name_column_index1 and merged_range.max_col >= name_column_index1:
                        name_range = merged_range
                        break
                if name_range:
                    for cell in sheet[name_range]:
                        name_route_map[cell.value].append(route)
                        name_route_ids[cell.value].append(f"{cell.value}: {route}")
            else:
                name_route_map[name].append(route)
                name_route_ids[name].append(f"{name}: {route}")

    return name_route_map, name_route_ids

##--------------------------------------------------------------------------------------------------- 

def create_report_document(name_route_ids, total_misloads, output_folder):
    def calculate_percentage(count, total):
        if total > 0:
            return round(count / total * 100, 2)
        else:
            return 0

    def omit_loader_name_from_route_ids(loader_name, route_ids):
        return [route_id.replace(loader_name, "").strip().replace(":", "") for route_id in route_ids]

    # Create a new Word document
    doc = Document()

    # Set the font to Arial and font size to 12 for the entire document
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(12)

    # Add the report title
    title = doc.add_heading("Three Rivers East: Daily Misload Report", level=1)
    title.bold = True
    title.runs[0].font.size = Pt(18)

    # Get the current date and weekday
    current_time = datetime.now()
    current_date = current_time.date()
    weekday = current_date.weekday()

    # Adjust the report date based on the weekday
    if weekday == 1:  # Tuesday
        previous_saturday = current_date - timedelta(days=(weekday + 2))
        report_date = previous_saturday.strftime("%Y-%m-%d")
    else:
        report_date = (current_time - timedelta(days=1)).strftime("%Y-%m-%d")

    # Add the report generation timestamp with bold font and font size 12
    p = doc.add_paragraph("Report Date: ", style='Normal')
    p.add_run(report_date).bold = True
    p.runs[1].font.size = Pt(12)

    # Add the total number of misloads with bold font and font size 12
    p = doc.add_paragraph("Total Misloads: ", style='Normal')
    p.add_run(str(total_misloads)).bold = True
    p.runs[1].font.size = Pt(12)

    file_path2 = find_most_recent_excel_file(folder_path2)
    if file_path2:
        workbook = load_workbook(file_path2)
        sheet = workbook.worksheets[sheet_index2]

        route_count_map = defaultdict(int)
        for column in sheet.iter_cols(min_col=route_column_index2 + 1, max_col=route_column_index2 + 1, min_row=1, values_only=True):
            for route in column:
                if route:
                    route_count_map[route] += 1

        # Create lists to store names and misload counts
        misload_list = []
        no_misload_list = []

        # Add the individual misloads to the appropriate list
        for name, routes in name_route_map.items():
            total_route_count = sum(route_count_map[route] for route in routes)
            if total_route_count > 0:
                misload_list.append((name, total_route_count))
            else:
                no_misload_list.append((name, total_route_count))

        # Sort the misload lists in descending order based on misload counts
        misload_list.sort(key=lambda x: x[1], reverse=True)
        no_misload_list.sort(key=lambda x: x[1], reverse=True)

        # Add the lists to the document in a grid format
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # Set the table header with bold font and font size 12
        header_cells = table.rows[0].cells
        header_cells[0].text = "Loader"
        header_cells[0].paragraphs[0].runs[0].bold = True
        header_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
        header_cells[1].text = "Misloads"
        header_cells[1].paragraphs[0].runs[0].bold = True
        header_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
        header_cells[2].text = "Percentage"
        header_cells[2].paragraphs[0].runs[0].bold = True
        header_cells[2].paragraphs[0].runs[0].font.size = Pt(12)
        header_cells[3].text = "Loader"
        header_cells[3].paragraphs[0].runs[0].bold = True
        header_cells[3].paragraphs[0].runs[0].font.size = Pt(12)
        header_cells[4].text = "Good Job!"
        header_cells[4].paragraphs[0].runs[0].bold = True
        header_cells[4].paragraphs[0].runs[0].font.size = Pt(12)

        # Add the data rows with background colors and percentage calculation
        for i in range(max(len(misload_list), len(no_misload_list))):
            cells = table.add_row().cells

            if i < len(misload_list):
                name, misload_count = misload_list[i]
                cells[0].text = name
                cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                cells[1].text = str(misload_count)
                cells[1].paragraphs[0].runs[0].font.size = Pt(12)
                percentage = calculate_percentage(misload_count, total_misloads)
                cells[2].text = str(percentage) + "%"
                cells[2].paragraphs[0].runs[0].font.size = Pt(12)
                set_cell_background_color(cells[1], misload_count)

            if i < len(no_misload_list):
                name, misload_count = no_misload_list[i]
                cells[3].text = name
                cells[3].paragraphs[0].runs[0].font.size = Pt(12)
                cells[4].text = str(misload_count)
                cells[4].paragraphs[0].runs[0].font.size = Pt(12)
                percentage = calculate_percentage(misload_count, total_misloads)
                cells[4].text = str(percentage) + "%"
                cells[4].paragraphs[0].runs[0].font.size = Pt(12)
                set_cell_background_color(cells[4], misload_count)

    # Add assigned load groups details on the first page
    doc.add_paragraph(style='Normal').add_run("\nAssigned Load Group Details").bold = True

    for name, route_ids in name_route_ids.items():
        route_ids_filtered = omit_loader_name_from_route_ids(name, route_ids)
        p = doc.add_paragraph(style='Normal')
        p.add_run("Loader: ").bold = True
        p.runs[0].font.size = Pt(9)
        p.add_run(name).font.size = Pt(9)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        if route_ids_filtered:
            p.add_run("\nRoute IDs: ").bold = True
            p.runs[2].font.size = Pt(9)
            p.add_run(", ".join(route_ids_filtered)).font.size = Pt(9)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    # Save the document with a timestamp in the file name
    output_file_name = f"Three_Rivers_East_Misloads_Report_{report_date}.docx"
    output_path = os.path.join(output_folder, output_file_name)
    doc.save(output_path)

    print(f"Report document saved at: {output_path}")
    
##---------------------------------------------------------------------------------------------------

def set_cell_background_color(cell, misload_count):
    if misload_count == 0:
        shading_element = parse_xml(r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls('w')))
    elif misload_count == 1:
        shading_element = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
    else:
        shading_element = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
    cell._element.append(shading_element)

##---------------------------------------------------------------------------------------------------

# Specify the folder paths and indices
folder_path1 = "/home/kbarakzai/Desktop/Staffing"
folder_path2 = "/home/kbarakzai/Desktop/Misloads"


sheet_index1 = 2
name_column_index1 = 17
route_column_index1 = 14
sheet_index2 = 0
route_column_index2 = 5

##---------------------------------------------------------------------------------------------------

# Find the most recent Excel files and extract the name-route associations from Staffing sheet
file_path1 = find_most_recent_excel_file(folder_path1)
if file_path1:
    name_route_map, name_route_ids = extract_name_route_associations(file_path1, sheet_index1, name_column_index1, route_column_index1)
    # Find most recent Excel File of Misloads sheet
    file_path2 = find_most_recent_excel_file(folder_path2)
    if file_path2:
        workbook = load_workbook(file_path2)
        sheet = workbook.worksheets[sheet_index2]

        route_count_map = defaultdict(int)
        for column in sheet.iter_cols(min_col=route_column_index2 + 1, max_col=route_column_index2 + 1, min_row=1, values_only=True):
            for route in column:
                if route:
                    route_count_map[route] += 1

        total_misloads = 0  # Variable to store the total number of misloads
        for name, routes in name_route_map.items():
            total_route_count = sum(route_count_map[route] for route in routes)
            total_misloads += total_route_count  # Add the misload count to the total
            print(f"{name}: {total_route_count}")

        print("\nTotal Misloads:", total_misloads)  # Print the total number of misloads

        # Specify the folder path for the output document
        output_folder = "/home/kbarakzai/Desktop/TR East Misload Reports"

        print("\n--Assigned Load Groups Details--")
        for name, route_ids in name_route_ids.items():
            route_ids_str = ', '.join(set(route_ids))
            print(f"{route_ids_str}")
        print()
        
        create_report_document(name_route_ids, total_misloads, output_folder)

##---------------------------------------------------------------------------------------------------


